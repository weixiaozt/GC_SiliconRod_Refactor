from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

doc = Document()

style = doc.styles['Normal']
style.font.name = '微软雅黑'
style.font.size = Pt(11)
style.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')

def h(text, level=1):
    p = doc.add_heading(text, level=level)
    for run in p.runs:
        run.font.name = '微软雅黑'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    return p

def p(text):
    para = doc.add_paragraph(text)
    return para

def table(headers, rows):
    t = doc.add_table(rows=1+len(rows), cols=len(headers))
    t.style = 'Light Grid Accent 1'
    for i, head in enumerate(headers):
        t.rows[0].cells[i].text = head
    for r, row in enumerate(rows, start=1):
        for i, val in enumerate(row):
            t.rows[r].cells[i].text = val
    return t

# ── 封面 ──
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title.add_run('SiRod Inspector\n用户使用手册')
run.font.size = Pt(28)
run.font.bold = True
run.font.name = '微软雅黑'
run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')

sub = doc.add_paragraph()
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = sub.add_run('光伏方棒隐裂检测系统')
r.font.size = Pt(16)
r.font.name = '微软雅黑'
r._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')

doc.add_paragraph()
doc.add_paragraph()

# 一、软件简介
h('一、软件简介', 1)
p('SiRod Inspector 是用于光伏方棒产线的隐裂/崩边在线检测上位机软件。'
  '它通过 TCP 接收视觉检测设备上传的结果与图像,提供实时展示、历史查询、缺陷图库、统计报表等功能,'
  '并自动将数据同步到 MySQL 数据库与飞书多维表格。')

# 二、运行环境
h('二、运行环境', 1)
table(['项目', '要求'], [
    ['操作系统', 'Windows 10 / 11'],
    ['显示分辨率', '推荐 1920×1080(最低 1280×720)'],
    ['数据库', 'MySQL 5.7+,库 b_xmartsql,表 squarstickresult'],
    ['网络', '飞书同步需外网;TCP 默认端口 3000'],
    ['磁盘', '图像保存路径(默认 D:/SiRod/images)需可写'],
])

# 三、启动软件
h('三、启动软件', 1)
h('方式一:免安装版(推荐现场使用)', 2)
p('进入 dist/SiRod_Inspector/ 目录,双击 SiRod_Inspector.exe。')
h('方式二:源码运行(开发调试)', 2)
p('pip install PyQt6 numpy Pillow pymysql matplotlib openpyxl requests')
p('python main.py')
p('启动后窗口自动最大化,右上角状态徽章显示"运行中"(绿色)即表示就绪。')

# 四、主界面布局
h('四、主界面布局', 1)
p('顶栏:Logo + 5 个页面导航 + 运行状态徽章 + 实时时间 + 产线编号')
p('底栏:三个设备状态指示灯(绿=正常,灰=未连接,红=异常) + 累计接收计数')

# 五、五大功能页面
h('五、五大功能页面', 1)

h('1. 总览页', 2)
p('实时展示最新一条检测数据(棒料编号、结果、缺陷类型、缺陷数),并统计当前班次的总数、OK 数、NG 数、'
  '良率与平均 CT 节拍。班次到时自动清零。')

h('2. 历史记录', 2)
p('按时间区间 / 结果(OK/NG) / 缺陷类型查询数据库历史记录,结果以表格显示,支持导出 Excel。')

h('3. 缺陷图库', 2)
p('只展示 NG 数据的图片卡片,顶部下拉框可按缺陷类型筛选。点击卡片可查看大图。新数据到达时自动新增卡片。')

h('4. 统计报表', 2)
p('按日/周/月汇总良率趋势、缺陷类型分布、班次对比等图表,可用于产线质量分析。')

h('5. 系统设置', 2)
p('分组配置所有参数,修改后点击保存即时写入 config.json:')
for item in [
    'TCP 通信:监听 IP / 端口(改后需重启)',
    '数据库:主机 / 端口 / 账号 / 密码 / 库表名',
    '飞书同步:启用开关 + AppID / Secret / AppToken / TableID',
    '图像存储:启用开关 + 根目录',
    '班次清零时间:支持多个时间点(如 08:00、20:00)',
    '缺陷类型:新增/删除,变更后图库筛选下拉同步更新',
    '产线编号:写入数据库与飞书的 line_id',
]:
    doc.add_paragraph(item, style='List Bullet')

# 六、典型使用流程
h('六、典型使用流程', 1)
for i, step in enumerate([
    '开机启动 — 双击 exe,等待主窗口出现。',
    '确认连接 — 检查底栏三个指示灯全部变绿;如有灰色,进入"系统设置"检查对应配置。',
    '正常运行 — 产线推送数据后,总览页实时刷新,底栏计数递增。',
    '查看 NG — 切换到"缺陷图库",按类型筛选查看。',
    '导出记录 — 在"历史记录"设置条件查询 → 点击导出 Excel。',
    '交接班 — 到达设置的清零时间后,总览数据自动归零开始新班统计。',
    '关闭软件 — 直接关闭窗口,程序会自动保存班次统计并断开所有连接。',
], start=1):
    doc.add_paragraph(f'{i}. {step}', style='List Number')

# 七、数据存放位置
h('七、数据存放位置', 1)
table(['类型', '位置'], [
    ['检测结果', 'MySQL b_xmartsql.squarstickresult + 飞书多维表'],
    ['图像文件', 'D:/SiRod/images/YYYY-MM-DD/{OK或NG}/{棒号}_{时分秒}.png'],
    ['班次统计缓存', '程序目录下 shift_stats.json'],
    ['运行日志', '程序目录下 logs/(保留 30 天)'],
])

# 八、常见问题
h('八、常见问题', 1)
table(['现象', '处理方法'], [
    ['启动报"模块导入失败"', 'pip install PyQt6 numpy Pillow pymysql matplotlib openpyxl requests'],
    ['底栏 TCP 红色', '端口被占用,到"系统设置"改端口后重启'],
    ['数据库指示灯灰色', '检查账号密码、MySQL 服务是否启动、库表是否创建'],
    ['飞书上传失败', '检查外网连接、AppID/Secret/AppToken 是否正确'],
    ['图像未保存', '确认"图像存储"已启用且目录可写'],
    ['班次未按时清零', '检查"班次清零时间"设置(定时器每 30 秒检查一次)'],
    ['程序异常', '查看 logs/ 下当日日志中的 ERROR/CRITICAL 关键字'],
])

# 九、技术支持
h('九、技术支持', 1)
p('如需排查问题,请提供:')
for item in [
    'logs/ 下的当日日志文件',
    'config.json(请隐去密码与飞书 Secret)',
    '异常发生时间点与现场截图',
]:
    doc.add_paragraph(item, style='List Bullet')

out = 'SiRod_Inspector_用户使用手册.docx'
doc.save(out)
print(f'已生成: {out}')
